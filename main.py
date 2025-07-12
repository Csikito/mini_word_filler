import datetime
import pandas as pd
from time import sleep
from docx import Document
from pathlib import Path


TEMPLATE_DIR = Path("templates")
OUTPUT_DIR = Path("output")
INPUT_FILE = Path("input.xlsx")

OUTPUT_DIR.mkdir(exist_ok=True)

timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
run_output_dir = OUTPUT_DIR / timestamp
run_output_dir.mkdir(parents=True, exist_ok=True)


def process_input():
    try:
        input_df = pd.read_excel(INPUT_FILE)
        return  dict(zip(input_df["key"], input_df["value"]))
    except Exception as e:
        print(e)
        return None

def find_and_replace_value(text, data):
    if "<<" and ">>" in text:
        for key, value in data.items():
            placeholder = f"<<{key}>>"
            if placeholder in text:
                value = value_formatter(value)
                text = text.replace(f'<<{key}>>', str(value))
    return text

def value_formatter(value):
    if isinstance(value, datetime.date):
        value = value.strftime("%Y-%m-%d")
    elif value == "tick":
        value = "🗹"
    elif value == "no_tick":
        value ="☐"
    return value

def fill_placeholders(doc: Document, data: dict):
    # Paragraph
    for p in doc.paragraphs:
        p.text = find_and_replace_value(p.text, data)
    # Table
    for tabla in doc.tables:
        for sor_tabla in tabla.rows:
            for cella in sor_tabla.cells:
                cella.text = find_and_replace_value(cella.text, data)
    return doc

def main():
    print("Start filling out files")
    print("")
    template_files = list(TEMPLATE_DIR.glob("*.docx"))
    data = process_input()
    generated_files = []

    if data and template_files:
        for template_path in template_files:
            try:
                doc = Document(template_path)
            except Exception as e:
                print(f"Bad file - {template_path} - {e}")
                print("")
                continue
            doc = fill_placeholders(doc, data)
            output_path = run_output_dir / f"{template_path.stem}_filled.docx"
            doc.save(output_path)
            generated_files.append(output_path.name)
            print(f"{output_path.name} file done!")
        print("")
        print("****************************")
        print("*** Files has been done! ***")
        print("****************************")


if __name__ == '__main__':
    main()
    sleep(3)