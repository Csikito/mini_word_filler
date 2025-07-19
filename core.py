import datetime
import pandas as pd
from docx import Document
from pathlib import Path
from tkinter import messagebox


def process_input(excel_path):
    try:
        input_df = pd.read_excel(excel_path)
        data = dict(zip(input_df["key"], input_df["value"]))
        return data
    except Exception as e:
        messagebox.showerror("Excel Error", f"Failed to read Excel file:\n{e}")
        return

def get_unique_output_folder(output_dir):
    selected_output_dir = output_dir or Path("output")
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    output_dir = selected_output_dir / timestamp
    output_dir.mkdir(parents=True, exist_ok=True)
    return output_dir

def find_and_replace_value(text, data):
    if "<<" and ">>" in text:
        for key, value in data.items():
            placeholder = f"<<{key}>>"
            if placeholder in text:
                value = value_formatter(value)
                text = text.replace(f'<<{key}>>', str(value))
    return text

def value_formatter(value):
    if isinstance(value, datetime.date):
        value = value.strftime("%Y-%m-%d")
    elif value == "tick":
        value = "🗹"
    elif value == "no_tick":
        value ="☐"
    return value

def fill_placeholders(doc: Document, data: dict):
    # Paragraph
    for p in doc.paragraphs:
        p.text = find_and_replace_value(p.text, data)
    # Table
    for tabla in doc.tables:
        for sor_tabla in tabla.rows:
            for cella in sor_tabla.cells:
                cella.text = find_and_replace_value(cella.text, data)
    return doc

def fill_files(data, template_files, output_dir):
    generated = 0

    for template in template_files:
        try:
            doc = Document(template)
            doc = fill_placeholders(doc, data)
            output_path = output_dir / f"{template.stem}_filled.docx"
            doc.save(output_path)
            generated += 1
        except Exception as e:
            messagebox.showerror("Processing Error", f"Error processing {template.name}:\n{e}")
    return generated
